import streamlit as st
import pandas as pd
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
import io

# ================= 1. 页面配置 =================
st.set_page_config(
    page_title="智能财务分析系统", 
    page_icon="📈",
    layout="wide"
)

# ================= 2. 核心工具函数 =================

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        if border_name in kwargs:
            edge = kwargs[border_name]
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), edge.get('val', 'single'))
            border.set(qn('w:sz'), str(edge.get('sz', 4)))
            border.set(qn('w:space'), str(edge.get('space', 0)))
            border.set(qn('w:color'), edge.get('color', 'auto'))
            tcBorders.append(border)

def create_word_table_file(df, title="数据表", bold_rows=None):
    """🔥 生成精排版 Word 表格 (审计底稿风格)"""
    doc = Document()
    
    # 设置页边距为窄边距
    section = doc.sections[0]
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体') 
        run.font.bold = True
        run.font.color.rgb = None

    export_df = df.reset_index()
    table = doc.add_table(rows=1, cols=len(export_df.columns))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False 
    
    # 动态计算列宽
    num_cols = len(export_df.columns)
    if num_cols > 5:
        first_col_w = Cm(5.0)
        other_col_w = Cm(2.2) 
    else:
        first_col_w = Cm(6.0)
        other_col_w = Cm(3.0)

    col_widths = [first_col_w] + [other_col_w] * (num_cols - 1)
    
    for i, width in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = width

    hdr_cells = table.rows[0].cells
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height = Cm(1.0)

    for i, col_name in enumerate(export_df.columns):
        cell = hdr_cells[i]
        cell.text = str(col_name)
        set_cell_border(cell, top={"val": "single", "sz": 12}, bottom={"val": "single", "sz": 12}, left={"val": "single", "sz": 4}, right={"val": "single", "sz": 4})
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10.5)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for r_idx, row in export_df.iterrows():
        row_cells = table.add_row().cells
        table.rows[r_idx+1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        table.rows[r_idx+1].height = Cm(0.6)
        
        subject_name = str(row[0]).strip()
        is_bold = False
        if bold_rows and subject_name in bold_rows: is_bold = True
        elif any(k in subject_name for k in ["合计", "总计", "净额", "净增加额", "构成"]): is_bold = True
        elif subject_name.endswith("：") or subject_name.endswith(":"): is_bold = True

        for i, val in enumerate(row):
            cell = row_cells[i]
            cell.text = str(val) if pd.notna(val) and val != "" else ""
            bottom_sz = 12 if r_idx == len(export_df) - 1 else 4
            set_cell_border(cell, top={"val": "single", "sz": 4}, bottom={"val": "single", "sz": bottom_sz}, left={"val": "single", "sz": 4}, right={"val": "single", "sz": 4})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            for run in paragraph.runs:
                run.font.size = Pt(10.5)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if is_bold: run.font.bold = True
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def extract_date_label(header_str):
    s = str(header_str).strip()
    match = re.search(r'[【\[](.*?)[】\]]', s)
    if match: return match.group(1)
    year = re.search(r'(\d{4})', s)
    if year: return f"{year.group(1)}年"
    return s

def fuzzy_load_excel(file_obj, sheet_name, header_row=None):
    try:
        xl = pd.ExcelFile(file_obj)
        all_sheet_names = xl.sheet_names
        target_sheet = None
        
        if sheet_name in all_sheet_names:
            target_sheet = sheet_name
        else:
            clean_target = sheet_name.replace(" ", "")
            for actual_name in all_sheet_names:
                if actual_name.replace(" ", "") == clean_target:
                    target_sheet = actual_name
                    break
        
        if target_sheet is None:
            return None, all_sheet_names

        if "财务指标" in sheet_name or "5-3" in sheet_name:
            return smart_load_ratios(file_obj, target_sheet)
        
        return pd.read_excel(file_obj, sheet_name=target_sheet, header=header_row), None

    except Exception as e:
        return None, [str(e)]

def smart_load_ratios(file_obj, sheet_name):
    try:
        df_raw = pd.read_excel(file_obj, sheet_name=sheet_name, header=None)
        header_idx = -1
        for i in range(10):
            row_values = df_raw.iloc[i].astype(str).values
            if any("项目" in v or "指标" in v for v in row_values):
                header_idx = i
                break
        if header_idx == -1: header_idx = 1
        df = pd.read_excel(file_obj, sheet_name=sheet_name, header=header_idx)
        cols = df.columns.tolist()
        date_col_indices = []
        for idx, col_name in enumerate(cols):
            s = str(col_name)
            if "年" in s or "T" in s or "202" in s or "期" in s:
                date_col_indices.append(idx)
        if len(date_col_indices) >= 3:
            target_cols = [0] + date_col_indices[:3]
        else:
            target_cols = [0, 2, 3, 4]
        df_final = df.iloc[:, target_cols]
        orig_cols = df_final.columns.tolist()
        d_labels = [extract_date_label(c) for c in orig_cols[1:]]
        df_final.columns = ['科目', 'T', 'T_1', 'T_2']
        df_final = df_final.dropna(subset=['科目'])
        df_final['科目'] = df_final['科目'].astype(str).str.strip()
        for c in ['T', 'T_1', 'T_2']:
            df_final[c] = pd.to_numeric(df_final[c], errors='coerce').fillna(0)
        df_final.set_index('科目', inplace=True)
        return df_final, d_labels
    except Exception as e:
        raise Exception(f"智能读取失败: {str(e)}")

def find_row_fuzzy(df, keywords, exclude_keywords=None, default_val=None):
    if isinstance(keywords, str): keywords = [keywords]
    clean_index = df.index.astype(str).str.replace(r'\s+', '', regex=True)
    found_rows = []
    for kw in keywords:
        clean_kw = kw.replace(" ", "")
        mask_exact = clean_index == clean_kw
        mask_contains = clean_index.str.contains(clean_kw, case=False, na=False)
        if exclude_keywords:
            for ex_kw in exclude_keywords:
                clean_ex = ex_kw.replace(" ", "")
                mask_contains = mask_contains & (~clean_index.str.contains(clean_ex, case=False, na=False))
        matched_indices = df.index[mask_exact | mask_contains].tolist()
        for idx in matched_indices:
            row = df.loc[idx]
            if isinstance(row, pd.DataFrame):
                for _, r in row.iterrows(): found_rows.append(r)
            else:
                found_rows.append(row)
    best_row = None
    max_non_zeros = -1
    for row in found_rows:
        non_zeros = 0
        if row['T'] != 0 and pd.notna(row['T']): non_zeros += 1
        if row['T_1'] != 0 and pd.notna(row['T_1']): non_zeros += 1
        if row['T_2'] != 0 and pd.notna(row['T_2']): non_zeros += 1
        if non_zeros > max_non_zeros:
            max_non_zeros = non_zeros
            best_row = row
    if best_row is not None: return best_row
    if default_val is not None: return default_val
    return pd.Series(0, index=df.columns)

# ================= 3. 业务逻辑处理函数 (Global) =================

def process_analysis_tab(df_raw, word_data_list, total_col_name, analysis_name, d_labels):
    """处理资产和负债结构分析"""
    try:
        total_row = find_row_fuzzy(df_raw, [total_col_name])
        df = df_raw.copy()
        # 过滤掉三年均为0的科目
        mask_keep = ~((df['T'] == 0) & (df['T_1'] == 0) & (df['T_2'] == 0)) 
        mask_title = df.index.astype(str).str.contains(r'[:：]')
        df = df[mask_keep | mask_title]

        for period in ['T', 'T_1', 'T_2']:
            total = total_row[period]
            df[f'占比_{period}'] = df[period] / total if total != 0 else 0.0

        tab1, tab2, tab3 = st.tabs(["📋 明细数据", "📝 综述文案", "📝 变动分析文案"])

        with tab1:
            display_df = pd.DataFrame(index=df.index)
            for p, label in zip(['T', 'T_1', 'T_2'], d_labels):
                display_df[label] = df[p].apply(lambda x: f"{x:,.2f}")
                display_df[f"{label}占比(%)"] = (df[f'占比_{p}'] * 100).apply(lambda x: f"{x:.2f}")
            
            # 清除标题行的数据显示
            for idx in display_df.index:
                if str(idx).strip().endswith(("：", ":")):
                    display_df.loc[idx] = ""
            
            st.dataframe(display_df, use_container_width=True)
            doc_file = create_word_table_file(display_df, title=f"{analysis_name}结构情况表")
            st.download_button(f"📥 下载 Word", doc_file, f"{analysis_name}明细.docx")

        with tab2:
            top_5 = df.sort_values(by='T', ascending=False).head(5).index.tolist()
            denom_text = "总资产" if analysis_name == "资产" else "负债总额"
            summary_text = f"在{denom_text}构成中，发行人{analysis_name}主要为 **{'、'.join(top_5)}** 等。"
            st.markdown(f"#### 📝 {analysis_name}综述文案")
            st.code(summary_text, language='text')

        with tab3:
            st.info(f"💡 **提示**：已根据数据生成科目变动分析文案草稿。")
            major_subjects = df[(df[f'占比_T'] > 0.01) & (~df.index.str.contains(r'合计|总计|总额'))].index.tolist()
            for subject in major_subjects:
                row = df.loc[subject]
                diff_curr = row['T'] - row['T_1']
                dir_curr = "增加" if diff_curr >= 0 else "减少"
                analysis_text = f"{d_labels[0]}末，发行人{subject}较{d_labels[1]}末{dir_curr}{abs(diff_curr):,.2f}万元。"
                with st.expander(f"📌 {subject}"):
                    st.code(analysis_text, language='text')
    except Exception as e:
        st.error(f"处理分析页面时出错: {e}")

def process_cash_flow_tab(df_raw, word_data_list, d_labels):
    """处理现金流量分析 (补全函数)"""
    st.subheader("现金流量表分析")
    st.dataframe(df_raw, use_container_width=True)
    doc_file = create_word_table_file(df_raw, title="现金流量分析表")
    st.download_button(f"📥 下载 Word", doc_file, "现金流分析.docx")

def process_profitability_tab(df_raw, word_data_list, d_labels):
    """处理盈利能力分析 (补全函数)"""
    st.subheader("盈利能力分析")
    st.dataframe(df_raw, use_container_width=True)
    doc_file = create_word_table_file(df_raw, title="盈利能力分析表")
    st.download_button(f"📥 下载 Word", doc_file, "盈利能力分析.docx")

def process_financial_ratios_tab(df_raw, word_data_list, d_labels):
    """处理财务指标分析 (补全函数)"""
    st.subheader("主要财务指标分析")
    st.dataframe(df_raw, use_container_width=True)
    doc_file = create_word_table_file(df_raw, title="主要财务指标表")
    st.download_button(f"📥 下载 Word", doc_file, "财务指标分析.docx")


# ================= 4. 侧边栏与状态 =================
if 'show_manual' not in st.session_state:
    st.session_state.show_manual = False

def go_to_manual():
    st.session_state.show_manual = True

def go_to_analysis():
    st.session_state.show_manual = False

with st.sidebar:
    st.title("🎛️ 操控台")
    analysis_page = st.radio(
        "请选择要生成的章节：", 
        ["(一) 资产结构分析", "(二) 负债结构分析", "(三) 现金流量分析", "(四) 财务指标分析", "(五) 盈利能力分析"],
        on_change=go_to_analysis 
    )
    st.markdown("---")
    uploaded_excel = st.file_uploader("Excel 底稿 (必须)", type=["xlsx", "xlsm"], on_change=go_to_analysis)
    st.markdown("---")
    if st.button("📘 使用说明书", use_container_width=True):
        go_to_manual()
        st.rerun()

# ================= 5. 主程序执行 =================

# 系统默认配置
DEFAULT_HEADER_ROW = 2 
SHEET_CONFIG = {
    "asset": "1.合并资产表",
    "liab": "2.合并负债及权益表",
    "profit": "3.合并利润表",
    "cash": "4.合并现金流量表",
    "ratios": "5-3主要财务指标计算-方案3（专用公司债）"
}

if not uploaded_excel or st.session_state.show_manual:
    st.title("📊 财务分析报告自动化助手")
    st.info("💡 请先在左侧上传符合标准审计底稿模版的 Excel 文件。")
    if not uploaded_excel:
        st.warning("👈 请先在左侧侧边栏上传 Excel 文件以开始使用。")
else:
    # 模拟空列表 (如果不需要RAG功能)
    word_data_list = [] 

    def get_clean_data(target_sheet_name):
        try:
            df, all_sheets_if_failed = fuzzy_load_excel(uploaded_excel, target_sheet_name, DEFAULT_HEADER_ROW)
            if df is None: return None, None, f"未找到 Sheet '{target_sheet_name}' (现有: {all_sheets_if_failed})"
            
            # 尝试截取前几列 
            df = df.iloc[:, [0, 4, 5, 6]]
            orig_cols = df.columns.tolist()
            d_labels = [extract_date_label(orig_cols[1]), extract_date_label(orig_cols[2]), extract_date_label(orig_cols[3])]
            df.columns = ['科目', 'T', 'T_1', 'T_2']
            df = df.dropna(subset=['科目'])
            df['科目'] = df['科目'].astype(str).str.strip()
            for c in ['T', 'T_1', 'T_2']:
                df[c] = pd.to_numeric(df[c], errors='coerce').fillna(0)
            df.set_index('科目', inplace=True)
            return df, d_labels, None
        except Exception as e: return None, None, str(e)

    st.header(f"📊 {analysis_page}")

    if analysis_page == "(一) 资产结构分析":
        df_asset, d_labels, err = get_clean_data(SHEET_CONFIG["asset"])
        if df_asset is not None: 
            process_analysis_tab(df_asset, word_data_list, "资产总计", "资产", d_labels)
        else: st.error(f"❌ 读取失败：{err}")

    elif analysis_page == "(二) 负债结构分析":
        df_liab, d_labels, err = get_clean_data(SHEET_CONFIG["liab"])
        if df_liab is not None:
            total_name = "负债合计" 
            if not find_row_fuzzy(df_liab, total_name).any() and find_row_fuzzy(df_liab, "负债总计").any():
                total_name = "负债总计"
            process_analysis_tab(df_liab, word_data_list, total_name, "负债", d_labels)
        else: st.error(f"❌ 读取失败：{err}")

    elif analysis_page == "(三) 现金流量分析":
        df_cash, d_labels, err = get_clean_data(SHEET_CONFIG["cash"])
        if df_cash is not None:
            process_cash_flow_tab(df_cash, word_data_list, d_labels)
        else: st.error(f"❌ 读取失败：{err}")

    elif analysis_page == "(四) 财务指标分析":
        df_ratios, d_labels = fuzzy_load_excel(uploaded_excel, SHEET_CONFIG["ratios"], DEFAULT_HEADER_ROW)
        if df_ratios is not None:
            process_financial_ratios_tab(df_ratios, word_data_list, d_labels)
        else: 
            st.error(f"❌ 读取失败：未找到 Sheet '{SHEET_CONFIG['ratios']}'")

    elif analysis_page == "(五) 盈利能力分析":
        df_profit, d_labels, err = get_clean_data(SHEET_CONFIG["profit"])
        if df_profit is not None:
            process_profitability_tab(df_profit, word_data_list, d_labels)
        else: st.error(f"❌ 读取失败：{err}")
